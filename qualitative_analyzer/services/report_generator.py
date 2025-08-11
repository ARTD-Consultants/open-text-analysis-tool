"""Report generation utilities for creating various output formats."""

import os
import logging
from datetime import datetime
from typing import Dict, Any, List, Optional
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.shared import OxmlElement, qn

from ..models.theme import ThemeHierarchy
from ..models.analysis_result import AnalysisSession

logger = logging.getLogger(__name__)


class ReportGenerator:
    """Generates comprehensive analysis reports in multiple formats."""
    
    def __init__(self, output_dir: str = "output"):
        """Initialize report generator."""
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        
        # Set up matplotlib style
        plt.style.use('default')
        sns.set_palette("viridis")
    
    def generate_comprehensive_report(
        self,
        session: AnalysisSession,
        theme_hierarchy: ThemeHierarchy,
        ai_client=None,
        report_title: str = "Qualitative Analysis Report"
    ) -> Dict[str, str]:
        """
        Generate a comprehensive report with multiple formats.
        
        Returns:
            Dictionary mapping format to file path
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        report_files = {}
        
        # Generate Excel summary
        excel_path = os.path.join(self.output_dir, f"analysis_summary_{timestamp}.xlsx")
        self._generate_excel_report(session, theme_hierarchy, excel_path)
        report_files["excel"] = excel_path
        
        # Generate Word report
        word_path = os.path.join(self.output_dir, f"detailed_report_{timestamp}.docx")
        self._generate_word_report(session, theme_hierarchy, word_path, report_title, ai_client)
        report_files["word"] = word_path
        
        # Generate visualizations
        chart_path = os.path.join(self.output_dir, f"theme_charts_{timestamp}.png")
        self._generate_theme_visualizations(theme_hierarchy, chart_path)
        report_files["charts"] = chart_path
        
        # Generate statistics summary
        stats_path = os.path.join(self.output_dir, f"statistics_{timestamp}.json")
        self._generate_statistics_file(session, theme_hierarchy, stats_path)
        report_files["statistics"] = stats_path
        
        logger.info(f"Generated comprehensive report files: {list(report_files.keys())}")
        return report_files
    
    def _generate_excel_report(
        self,
        session: AnalysisSession,
        theme_hierarchy: ThemeHierarchy,
        file_path: str
    ) -> None:
        """Generate Excel report with multiple sheets."""
        with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
            
            # Sheet 1: Analysis Results with Theme Tracking
            results_data = []
            for result in session.get_all_results():
                results_data.append({
                    "Entry ID": result.entry_id,
                    "Original Text": result.original_text,
                    "Summary": result.summary,
                    "Original Themes (AI Extracted)": ", ".join(result.original_themes),
                    "Consolidated Themes (Final)": ", ".join(result.themes),
                    "Original Theme Count": len(result.original_themes),
                    "Final Theme Count": len(result.themes),
                    "Themes Filtered": len(result.original_themes) - len([t for t in result.themes if t != "NA"]),
                    "Contains NA": "NA" in result.themes,
                    "Average Confidence": result.get_average_confidence(),
                    "Sentiment Score": result.sentiment_score,
                    "Sentiment Label": result.sentiment_label,
                    "Processing Time": result.processing_timestamp.isoformat()
                })
            
            results_df = pd.DataFrame(results_data)
            results_df.to_excel(writer, sheet_name='Analysis Results', index=False)
            
            # Sheet 2: Theme Summary
            theme_data = []
            all_themes = theme_hierarchy.get_all_themes_flat()
            logger.info(f"Found {len(all_themes)} themes for report")
            
            for theme in all_themes:
                try:
                    parent_name = "Root"
                    if theme.parent_id:
                        parent_theme = theme_hierarchy.get_parent(theme.theme_id)
                        parent_name = parent_theme.name if parent_theme else "Unknown Parent"
                    
                    theme_data.append({
                        "Theme Name": theme.name,
                        "Frequency": getattr(theme, 'frequency', 0),
                        "Average Confidence": getattr(theme, 'average_confidence', 0.0),
                        "Parent Theme": parent_name,
                        "Children Count": len(getattr(theme, 'children_ids', [])),
                        "Hierarchy Path": " > ".join(theme_hierarchy.get_theme_path(theme.theme_id))
                    })
                except Exception as theme_error:
                    logger.warning(f"Error processing theme {theme.name}: {theme_error}")
                    # Add minimal theme data
                    theme_data.append({
                        "Theme Name": theme.name,
                        "Frequency": 0,
                        "Average Confidence": 0.0,
                        "Parent Theme": "Root",
                        "Children Count": 0,
                        "Hierarchy Path": theme.name
                    })
            
            logger.info(f"Created theme data with {len(theme_data)} entries")
            
            if theme_data:
                theme_df = pd.DataFrame(theme_data)
                logger.info(f"Theme DataFrame columns: {list(theme_df.columns)}")
                if 'Frequency' in theme_df.columns:
                    theme_df = theme_df.sort_values('Frequency', ascending=False)
                theme_df.to_excel(writer, sheet_name='Theme Summary', index=False)
            else:
                # Create empty theme summary sheet
                empty_df = pd.DataFrame({"Message": ["No themes found"]})
                empty_df.to_excel(writer, sheet_name='Theme Summary', index=False)
            
            # Sheet 3: Session Statistics
            stats = session.get_session_statistics()
            stats_data = [{"Metric": k, "Value": v} for k, v in stats.items()]
            stats_df = pd.DataFrame(stats_data)
            stats_df.to_excel(writer, sheet_name='Session Statistics', index=False)
            
            # Sheet 4: Theme Co-occurrence
            cooccurrence_data = self._calculate_theme_cooccurrence(session.get_all_results())
            if cooccurrence_data:
                cooccurrence_df = pd.DataFrame(cooccurrence_data)
                cooccurrence_df.to_excel(writer, sheet_name='Theme Co-occurrence', index=False)
        
        logger.info(f"Excel report saved to {file_path}")
    
    def _generate_word_report(
        self,
        session: AnalysisSession,
        theme_hierarchy: ThemeHierarchy,
        file_path: str,
        title: str,
        ai_client=None
    ) -> None:
        """Generate detailed Word report."""
        doc = Document()
        
        # Title and metadata
        title_paragraph = doc.add_heading(title, level=0)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata table
        doc.add_heading('Analysis Overview', level=1)
        stats = session.get_session_statistics()
        
        overview_data = [
            ["Analysis Date", datetime.now().strftime("%B %d, %Y")],
            ["Total Entries Analyzed", stats["total_entries"]],
            ["Unique Themes Identified", stats["unique_themes"]],
            ["Processing Duration", f"{stats.get('session_duration', 0):.1f} seconds"],
            ["Average Themes per Entry", f"{stats['average_themes_per_entry']:.1f}"],
            ["Overall Confidence", f"{stats['overall_confidence']:.1%}"]
        ]
        
        table = doc.add_table(rows=len(overview_data), cols=2)
        table.style = 'Table Grid'
        
        for i, (key, value) in enumerate(overview_data):
            table.cell(i, 0).text = key
            table.cell(i, 1).text = str(value)
        
        # Theme frequency chart
        doc.add_heading('Theme Frequency Overview', level=1)
        chart_path = self._create_theme_frequency_chart(theme_hierarchy)
        doc.add_picture(chart_path, width=Inches(6.0))
        
        # Detailed theme analysis
        doc.add_heading('Detailed Theme Analysis', level=1)
        
        top_themes = theme_hierarchy.get_all_themes_flat()[:10]
        
        for i, theme in enumerate(top_themes):
            # Theme heading
            doc.add_heading(f"{i+1}. {theme.name}", level=2)
            
            # Theme statistics with proper formatting
            freq_p = doc.add_paragraph()
            freq_p.add_run("Frequency: ").bold = True
            freq_p.add_run(f"{theme.frequency} occurrences ({theme.frequency/stats['total_entries']*100:.1f}% of entries)")
            
            conf_p = doc.add_paragraph()
            conf_p.add_run("Average Confidence: ").bold = True
            conf_p.add_run(f"{theme.average_confidence:.1%}")
            
            if theme.parent_id:
                parent = theme_hierarchy.get_parent(theme.theme_id)
                if parent:
                    parent_p = doc.add_paragraph()
                    parent_p.add_run("Parent Theme: ").bold = True
                    parent_p.add_run(parent.name)
            
            # Get example texts for this theme
            theme_examples = self._get_theme_examples(session.get_all_results(), theme.name, max_examples=5)
            
            if ai_client and theme_examples:
                # Generate AI analysis
                try:
                    analysis = ai_client.generate_theme_analysis(
                        theme.name,
                        theme_examples,
                        theme.frequency,
                        max_words=300
                    )
                    doc.add_heading("Analysis", level=3)
                    doc.add_paragraph(analysis)
                except Exception as e:
                    logger.warning(f"Failed to generate AI analysis for theme {theme.name}: {e}")
            
            # Add representative quotes
            if theme_examples:
                doc.add_heading("Representative Quotes", level=3)
                for j, example in enumerate(theme_examples[:3], 1):
                    p = doc.add_paragraph()
                    p.add_run(f"Quote {j}: ").bold = True
                    p.add_run(f'"{example}"')
            
            # Add page break except for last theme
            if i < len(top_themes) - 1:
                doc.add_page_break()
        
        doc.save(file_path)
        logger.info(f"Word report saved to {file_path}")
    
    def _generate_theme_visualizations(
        self,
        theme_hierarchy: ThemeHierarchy,
        file_path: str
    ) -> None:
        """Generate comprehensive theme visualizations."""
        fig, axes = plt.subplots(2, 2, figsize=(15, 12))
        fig.suptitle('Theme Analysis Visualizations', fontsize=16, fontweight='bold')
        
        themes = theme_hierarchy.get_all_themes_flat()[:15]  # Top 15 themes
        theme_names = [theme.name for theme in themes]
        frequencies = [theme.frequency for theme in themes]
        confidences = [theme.average_confidence for theme in themes]
        
        # 1. Horizontal bar chart of theme frequencies
        ax1 = axes[0, 0]
        bars = ax1.barh(theme_names, frequencies, color=sns.color_palette("viridis", len(themes)))
        ax1.set_title('Theme Frequencies')
        ax1.set_xlabel('Frequency')
        
        # Add frequency labels
        for bar in bars:
            width = bar.get_width()
            ax1.text(width + 0.1, bar.get_y() + bar.get_height()/2, 
                    f'{int(width)}', ha='left', va='center')
        
        # 2. Theme confidence scores
        ax2 = axes[0, 1]
        ax2.scatter(frequencies, confidences, s=100, alpha=0.6, color='coral')
        ax2.set_title('Theme Confidence vs Frequency')
        ax2.set_xlabel('Frequency')
        ax2.set_ylabel('Average Confidence')
        
        # Add theme labels to points
        for i, name in enumerate(theme_names):
            ax2.annotate(name, (frequencies[i], confidences[i]), 
                        xytext=(5, 5), textcoords='offset points', 
                        fontsize=8, alpha=0.7)
        
        # 3. Theme distribution pie chart (top 10)
        ax3 = axes[1, 0]
        top_10_themes = themes[:10]
        top_10_names = [theme.name for theme in top_10_themes]
        top_10_freqs = [theme.frequency for theme in top_10_themes]
        
        # Add "Others" category if there are more themes
        if len(themes) > 10:
            others_freq = sum(theme.frequency for theme in themes[10:])
            top_10_names.append('Others')
            top_10_freqs.append(others_freq)
        
        ax3.pie(top_10_freqs, labels=top_10_names, autopct='%1.1f%%', startangle=90)
        ax3.set_title('Theme Distribution')
        
        # 4. Confidence distribution histogram
        ax4 = axes[1, 1]
        ax4.hist(confidences, bins=10, alpha=0.7, color='skyblue', edgecolor='black')
        ax4.set_title('Theme Confidence Distribution')
        ax4.set_xlabel('Average Confidence')
        ax4.set_ylabel('Number of Themes')
        
        plt.tight_layout()
        plt.savefig(file_path, dpi=300, bbox_inches='tight')
        plt.close()
        
        logger.info(f"Theme visualizations saved to {file_path}")
    
    def _create_theme_frequency_chart(self, theme_hierarchy: ThemeHierarchy) -> str:
        """Create a standalone theme frequency chart."""
        themes = theme_hierarchy.get_all_themes_flat()[:15]
        theme_names = [theme.name for theme in themes]
        frequencies = [theme.frequency for theme in themes]
        
        plt.figure(figsize=(10, 6))
        bars = plt.barh(theme_names, frequencies, color=sns.color_palette("viridis", len(themes)))
        
        # Add frequency labels
        for bar in bars:
            width = bar.get_width()
            plt.text(width + 0.3, bar.get_y() + bar.get_height()/2, 
                    f'{int(width)}', ha='left', va='center')
        
        plt.title('Theme Frequency')
        plt.xlabel('Frequency')
        plt.tight_layout()
        
        chart_path = os.path.join(self.output_dir, "temp_theme_chart.png")
        plt.savefig(chart_path, dpi=300, bbox_inches='tight')
        plt.close()
        
        return chart_path
    
    def _generate_statistics_file(
        self,
        session: AnalysisSession,
        theme_hierarchy: ThemeHierarchy,
        file_path: str
    ) -> None:
        """Generate detailed statistics JSON file."""
        import json
        
        stats = {
            "session_info": session.get_session_statistics(),
            "theme_hierarchy_stats": theme_hierarchy.get_statistics(),
            "theme_details": {
                theme.name: {
                    "frequency": theme.frequency,
                    "average_confidence": theme.average_confidence,
                    "hierarchy_level": len(theme_hierarchy.get_theme_path(theme.theme_id)),
                    "has_children": len(theme.children_ids) > 0
                }
                for theme in theme_hierarchy.get_all_themes_flat()
            },
            "processing_summary": {
                batch.batch_id: batch.to_dict()
                for batch in session.batches
            }
        }
        
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(stats, f, indent=2, default=str)
        
        logger.info(f"Statistics file saved to {file_path}")
    
    def _calculate_theme_cooccurrence(self, results: List) -> List[Dict[str, Any]]:
        """Calculate theme co-occurrence statistics."""
        cooccurrence_data = []
        theme_pairs = {}
        
        for result in results:
            themes = result.themes
            # Generate all pairs of themes for this entry
            for i in range(len(themes)):
                for j in range(i + 1, len(themes)):
                    pair = tuple(sorted([themes[i], themes[j]]))
                    theme_pairs[pair] = theme_pairs.get(pair, 0) + 1
        
        # Convert to list format
        for (theme1, theme2), count in theme_pairs.items():
            cooccurrence_data.append({
                "Theme 1": theme1,
                "Theme 2": theme2,
                "Co-occurrence Count": count,
                "Percentage": round(count / len(results) * 100, 2)
            })
        
        # Sort by co-occurrence count
        cooccurrence_data.sort(key=lambda x: x["Co-occurrence Count"], reverse=True)
        
        return cooccurrence_data[:50]  # Top 50 co-occurrences
    
    def _get_theme_examples(
        self,
        results: List,
        theme_name: str,
        max_examples: int = 5
    ) -> List[str]:
        """Get example texts for a specific theme."""
        examples = []
        
        for result in results:
            if theme_name in result.themes and len(examples) < max_examples:
                examples.append(result.original_text)
        
        return examples
    
    def generate_markdown_report(
        self,
        session: AnalysisSession,
        theme_hierarchy: ThemeHierarchy,
        file_path: str
    ) -> None:
        """Generate a markdown format report."""
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write("# Qualitative Analysis Report\\n\\n")
            f.write(f"**Generated:** {datetime.now().strftime('%B %d, %Y at %I:%M %p')}\\n\\n")
            
            # Overview
            stats = session.get_session_statistics()
            f.write("## Analysis Overview\\n\\n")
            f.write(f"- **Total Entries:** {stats['total_entries']}\\n")
            f.write(f"- **Unique Themes:** {stats['unique_themes']}\\n")
            f.write(f"- **Average Themes per Entry:** {stats['average_themes_per_entry']:.1f}\\n")
            f.write(f"- **Overall Confidence:** {stats['overall_confidence']:.1%}\\n\\n")
            
            # Top themes
            f.write("## Top Themes\\n\\n")
            f.write("| Rank | Theme | Frequency | Avg Confidence |\\n")
            f.write("|------|--------|-----------|----------------|\\n")
            
            for i, theme in enumerate(theme_hierarchy.get_all_themes_flat()[:15], 1):
                f.write(f"| {i} | {theme.name} | {theme.frequency} | {theme.average_confidence:.1%} |\\n")
            
            f.write("\\n")
        
        logger.info(f"Markdown report saved to {file_path}")